from langchain_core.tools import tool
import pandas as pd
from docx import Document
import os
import json
from pypdf import PdfReader
import ezdxf
from ezdxf import audit
import math

@tool
def search_standards(query: str) -> str:
    """Tra cứu Tiêu chuẩn thiết kế MEPF (TCVN, ASHRAE, NFPA...) từ cơ sở dữ liệu nội bộ."""
    print(f"\n[Tool] Tra cứu tiêu chuẩn thực: {query}")
    try:
        from langchain_openai import OpenAIEmbeddings
        from langchain_community.vectorstores import FAISS
        from src.config import settings
        import os
        
        api_key = settings.openai_api_key or os.getenv("OPENAI_API_KEY", "dummy_key_to_prevent_crash_on_import")
        embeddings = OpenAIEmbeddings(api_key=api_key)
        
        index_path = "faiss_index"
        if not os.path.exists(index_path):
            return "Hệ thống RAG chưa được khởi tạo. Vui lòng thêm tài liệu vào 'data/standards/' và chạy 'uv run python src/ingest.py'."
            
        vectorstore = FAISS.load_local(index_path, embeddings, allow_dangerous_deserialization=True)
        docs = vectorstore.similarity_search(query, k=3)
        
        if not docs:
            return "Không tìm thấy thông tin tiêu chuẩn nào khớp với yêu cầu."
            
        result = f"Kết quả RAG Tiêu chuẩn cho '{query}':\n"
        for i, doc in enumerate(docs, 1):
            source = doc.metadata.get('source', 'Unknown')
            result += f"\n--- Trích đoạn {i} (Nguồn: {source}) ---\n"
            result += doc.page_content + "\n"
            
        return result
    except Exception as e:
        return f"Lỗi tra cứu tiêu chuẩn RAG: {e}"

@tool
def search_web(query: str) -> str:
    """Tìm kiếm thông tin trên internet."""
    print(f"\n[Tool] Searching web for: {query}")
    return f"Kết quả mô phỏng cho '{query}': Tìm thấy nhiều tài liệu liên quan."

@tool
def calculate(expression: str) -> str:
    """Thực hiện tính toán toán học cơ bản (ví dụ: '25 * 4')."""
    print(f"\n[Tool] Calculating: {expression}")
    try:
        result = eval(expression, {"__builtins__": {}})
        return f"Kết quả: {result}"
    except Exception as e:
        return f"Lỗi tính toán: {e}"

@tool
def list_directory(path: str = ".") -> str:
    """Liệt kê danh sách các file trong thư mục để xem có file nào tồn tại."""
    print(f"\n[Tool] Listing directory: {path}")
    try:
        files = os.listdir(path)
        return f"Files trong '{path}': {', '.join(files)}"
    except Exception as e:
        return f"Lỗi đọc thư mục: {e}"

@tool
def read_excel(file_path: str) -> str:
    """Đọc nội dung từ file Excel (.xlsx)."""
    print(f"\n[Tool] Reading Excel: {file_path}")
    try:
        df = pd.read_excel(file_path)
        return f"Dữ liệu Excel:\n{df.to_string(index=False)}"
    except Exception as e:
        return f"Lỗi đọc Excel: {e}"

@tool
def write_excel(file_path: str, json_data: str) -> str:
    """Tạo hoặc ghi file Excel (.xlsx). json_data là danh sách các object dưới dạng chuỗi JSON đại diện cho các dòng. Ví dụ: '[{"STT": 1, "Vật tư": "Ống", "KL": 10}]'"""
    print(f"\n[Tool] Writing Excel: {file_path}")
    try:
        if not file_path.endswith('.xlsx'):
            file_path += '.xlsx'
            
        dir_name = os.path.dirname(file_path)
        if dir_name:
            os.makedirs(dir_name, exist_ok=True)
            
        data = json.loads(json_data)
        df = pd.DataFrame(data)
        df.to_excel(file_path, index=False)
        return f"Đã ghi đè/tạo thành công file Excel tại: {file_path}"
    except Exception as e:
        return f"Lỗi ghi Excel: {e}"

@tool
def read_word(file_path: str) -> str:
    """Đọc nội dung từ file Word (.docx)."""
    print(f"\n[Tool] Reading Word: {file_path}")
    try:
        doc = Document(file_path)
        full_text = [para.text for para in doc.paragraphs]
        return "\n".join(full_text)
    except Exception as e:
        return f"Lỗi đọc Word: {e}"

@tool
def write_word(file_path: str, content: str, font_name: str = 'Arial') -> str:
    """Tạo hoặc ghi file Word (.docx) với nội dung được truyền vào. Tham số font_name hỗ trợ 'Arial' hoặc 'Times New Roman'."""
    print(f"\n[Tool] Writing Word: {file_path}")
    try:
        from docx.shared import Pt
        doc = Document()
        # Thiết lập Font chữ chuẩn Unicode (Arial / Times New Roman) cho tiếng Việt
        style = doc.styles['Normal']
        font = style.font
        if font_name not in ['Arial', 'Times New Roman']:
            font_name = 'Arial'
        font.name = font_name
        font.size = Pt(12)
        
        doc.add_paragraph(content)
        doc.save(file_path)
        return f"Đã lưu nội dung vào file Word tại: {file_path} (Font: {font_name})"
    except Exception as e:
        return f"Lỗi ghi Word: {e}"

@tool
def read_pdf(file_path: str) -> str:
    """Đọc và trích xuất toàn bộ văn bản từ file PDF."""
    print(f"\n[Tool] Reading PDF: {file_path}")
    try:
        reader = PdfReader(file_path)
        text = ""
        for page in reader.pages:
            text += page.extract_text() + "\n"
        return f"Nội dung PDF ({len(reader.pages)} trang):\n{text[:5000]}..." 
    except Exception as e:
        return f"Lỗi đọc PDF: {e}"

@tool
def read_cad(file_path: str) -> str:
    """Đọc file CAD (.dxf) và trả về thống kê thư viện block, block attributes, chiều dài, và layer sau khi đã làm sạch."""
    print(f"\n[Tool] Reading, Cleaning & Extracting CAD: {file_path}")
    try:
        doc = ezdxf.readfile(file_path)
        
        auditor = audit.Auditor(doc)
        auditor.run()
        audit_fixes = len(auditor.fixes)
        
        block_defs = []
        for block in doc.blocks:
            if not block.is_layout_block and not block.name.startswith('*'):
                block_defs.append(block.name)
                
        msp = doc.modelspace()
        layer_counts = {}
        block_instances = []
        layer_lengths = {}
        
        for entity in msp:
            layer = entity.dxf.layer
            layer_counts[layer] = layer_counts.get(layer, 0) + 1
            dxftype = entity.dxftype()
            
            if dxftype == 'LINE':
                start = entity.dxf.start
                end = entity.dxf.end
                dist = math.hypot(end.x - start.x, end.y - start.y)
                layer_lengths[layer] = layer_lengths.get(layer, 0.0) + dist
            
            if dxftype == 'INSERT':
                b_name = entity.dxf.name
                attribs = {}
                if entity.has_attribs:
                    for attrib in entity.attribs:
                        attribs[attrib.dxf.tag] = attrib.dxf.text
                block_instances.append({"name": b_name, "attribs": attribs})
                
        block_summary = {}
        for b in block_instances:
            b_name = b['name']
            attr_str = json.dumps(b['attribs'], ensure_ascii=False) if b['attribs'] else "No Attributes"
            key = f"{b_name} | Thuộc tính: {attr_str}"
            block_summary[key] = block_summary.get(key, 0) + 1
            
        result = f"Đã làm sạch (Audit). Sửa {audit_fixes} lỗi.\n\n"
        result += f"THƯ VIỆN BLOCK CÓ SẴN (Definitions): {', '.join(block_defs) if block_defs else 'Không có'}\n\n"
        
        result += "THỐNG KÊ LAYER TRÊN MODELSPACE:\n"
        for k, v in layer_counts.items():
            l_info = f"- Layer '{k}': {v} đối tượng"
            if k in layer_lengths and layer_lengths[k] > 0:
                l_info += f" (Tổng dài Line tham khảo: {layer_lengths[k]:.2f})"
            result += l_info + "\n"
            
        result += "\nTHỐNG KÊ BLOCK THỰC TẾ & THUỘC TÍNH (Attributes):\n"
        if not block_summary:
            result += "(Không có block nào)\n"
        else:
            for k, v in block_summary.items():
                result += f"- Block: {k} -> Số lượng: {v}\n"
                
        return result
    except Exception as e:
        return f"Lỗi xử lý CAD (.dxf): {e}"

@tool
def write_cad(file_path: str, layers: str) -> str:
    """Tạo một file CAD mới (.dxf) sạch sẽ với các layer định trước. Tham số layers: chuỗi ngăn cách bởi dấu phẩy."""
    print(f"\n[Tool] Writing CAD: {file_path}")
    try:
        doc = ezdxf.new('R2010')
        layer_list = [l.strip() for l in layers.split(',') if l.strip()]
        for layer in layer_list:
            doc.layers.add(name=layer)
            
        doc.saveas(file_path)
        return f"Đã tạo thành công bản vẽ CAD tại {file_path} với các layers: {', '.join(layer_list)}"
    except Exception as e:
        return f"Lỗi tạo CAD (.dxf): {e}"

import sys
from io import StringIO

@tool
def execute_python_code(code: str) -> str:
    """
    Thực thi mã Python động. 
    Được dùng để Họa viên CAD tự viết code ezdxf vẽ Block mới và lưu vào 'data/blocks/mepf_library.dxf'.
    """
    print("\n[Tool] Executing Custom Python Code")
    try:
        old_stdout = sys.stdout
        redirected_output = sys.stdout = StringIO()
        
        local_env = {}
        exec(code, globals(), local_env)
        
        sys.stdout = old_stdout
        return f"Thực thi Python thành công. Output:\n{redirected_output.getvalue()}"
    except Exception as e:
        sys.stdout = old_stdout
        return f"Lỗi quá trình thực thi Python: {e}"

@tool
def ai_block_recovery(file_path: str, layer: str, shape: str, dimensions: str, replacement_block: str) -> str:
    """Khôi phục các thiết bị bị phá vỡ (exploded) thành Block chuẩn.
    - layer: Tên layer chứa các nét vẽ rời rạc.
    - shape: 'circle' (hình tròn) hoặc 'rectangle' (hình chữ nhật).
    - dimensions: Với circle là 'bán kính' (ví dụ: '100'). Với rectangle là 'dài,rộng' (ví dụ '600,600').
    - replacement_block: Tên Block mới sẽ được chèn vào.
    """
    print(f"\n[Tool] AI Block Recovery: {file_path}, Layer={layer}, Shape={shape}")
    try:
        from ezdxf.addons import importer
        import os
        
        if not os.path.exists(file_path):
            return f"Lỗi: Không tìm thấy file {file_path}"
            
        doc = ezdxf.readfile(file_path)
        msp = doc.modelspace()
        
        library_path = os.path.join("data", "blocks", "mepf_library.dxf")
        lib_doc = None
        if os.path.exists(library_path):
            lib_doc = ezdxf.readfile(library_path)
            
        if replacement_block not in doc.blocks and lib_doc and replacement_block in lib_doc.blocks:
            imp = importer.Importer(lib_doc, doc)
            imp.import_block(replacement_block)
            imp.finalize()
            
        if replacement_block not in doc.blocks:
            return f"Lỗi: Block '{replacement_block}' không tồn tại trong Thư viện Tổng kho."
            
        centers = []
        entities_to_delete = []
        max_dim = 0
        
        if shape.lower() == "circle":
            target_r = float(dimensions)
            max_dim = target_r * 2
            for entity in msp.query(f'CIRCLE[layer=="{layer}"]'):
                r = entity.dxf.radius
                if abs(r - target_r) / target_r <= 0.05:
                    centers.append((entity.dxf.center.x, entity.dxf.center.y))
                    
        elif shape.lower() == "rectangle":
            dims = dimensions.split(",")
            if len(dims) == 2:
                target_w, target_h = float(dims[0]), float(dims[1])
                max_dim = max(target_w, target_h)
                target_area = target_w * target_h
                for entity in msp.query(f'LWPOLYLINE[layer=="{layer}"]'):
                    if entity.closed or len(entity) >= 4:
                        points = entity.get_points()
                        xs = [p[0] for p in points]
                        ys = [p[1] for p in points]
                        w = max(xs) - min(xs)
                        h = max(ys) - min(ys)
                        area = w * h
                        if area > 0 and abs(area - target_area) / target_area <= 0.1:
                            cx = (max(xs) + min(xs)) / 2
                            cy = (max(ys) + min(ys)) / 2
                            centers.append((cx, cy))
                            
        # Dọn rác chuyên sâu: Quét và xóa MỌI nét vẽ (LINE, PLINE, TEXT) nằm lọt thỏm trong vùng Block
        if max_dim > 0 and centers:
            tolerance = max_dim * 0.6  # Phạm vi dọn rác (Bao trùm block + 10% an toàn)
            for entity in msp.query(f'*[layer=="{layer}"]'):
                px, py = None, None
                if hasattr(entity.dxf, 'start'):
                    px, py = entity.dxf.start.x, entity.dxf.start.y
                elif hasattr(entity.dxf, 'center'):
                    px, py = entity.dxf.center.x, entity.dxf.center.y
                elif hasattr(entity.dxf, 'insert'):
                    px, py = entity.dxf.insert.x, entity.dxf.insert.y
                elif entity.dxftype() == 'LWPOLYLINE':
                    try:
                        pts = entity.get_points()
                        px, py = pts[0][0], pts[0][1]
                    except:
                        pass
                
                if px is not None and py is not None:
                    for cx, cy in centers:
                        if abs(px - cx) <= tolerance and abs(py - cy) <= tolerance:
                            entities_to_delete.append(entity)
                            break
                            
        for e in set(entities_to_delete):
            try:
                msp.delete_entity(e)
            except:
                pass
            
        for cx, cy in centers:
            msp.add_blockref(replacement_block, (cx, cy), dxfattribs={'layer': layer})
            
        doc.saveas(file_path)
        return f"AI Recovery thành công: Đã tìm thấy và phục hồi {len(centers)} đối tượng '{shape}' thành Block '{replacement_block}'."
    except Exception as e:
        return f"Lỗi phục hồi Block: {e}"

@tool
def edit_cad(file_path: str, actions_json: str) -> str:
    """Chỉnh sửa file CAD (.dxf) hiện tại. Luôn Audit làm sạch file trước.
    actions_json là danh sách các dict. Ví dụ:
    - Thêm layer: {"action": "add_layer", "name": "MEP_DIEN"}
    - Thêm text: {"action": "add_text", "text": "Phong Khach", "x": 0, "y": 0, "layer": "MEP_DIEN", "font_name": "Times New Roman"}
    - Chèn block: {"action": "insert_block", "name": "TU_DIEN", "x": 10, "y": 10, "layer": "MEP_DIEN", "scale": 1.0, "rotation": 0}
    - Đồng bộ font (chống lỗi tiếng Việt): {"action": "fix_fonts", "font_name": "Arial"}
    """
    print(f"\n[Tool] Editing CAD: {file_path}")
    try:
        if not os.path.exists(file_path):
            return f"Lỗi: Không tìm thấy file {file_path}"
            
        doc = ezdxf.readfile(file_path)
        msp = doc.modelspace()
        
        auditor = audit.Auditor(doc)
        auditor.run()
        audit_fixes = len(auditor.fixes)
        
        actions = json.loads(actions_json)
        results = []
        
        # Tải Master Library (Tổng kho Block)
        from ezdxf.addons import importer
        library_path = os.path.join("data", "blocks", "mepf_library.dxf")
        lib_doc = None
        if os.path.exists(library_path):
            lib_doc = ezdxf.readfile(library_path)
            
        # Khởi tạo Style chữ chuẩn Unicode để tránh lỗi font tiếng Việt trong CAD
        if 'VIETNAMESE_ARIAL' not in doc.styles:
            doc.styles.new('VIETNAMESE_ARIAL', dxfattribs={'font': 'arial.ttf'})
        if 'VIETNAMESE_TIMES' not in doc.styles:
            doc.styles.new('VIETNAMESE_TIMES', dxfattribs={'font': 'times.ttf'})
            
        for act in actions:
            action_type = act.get("action")
            if action_type == "fix_fonts":
                font_name = act.get("font_name", "Arial")
                ttf_file = "times.ttf" if font_name == "Times New Roman" else "arial.ttf"
                # Đổi font của toàn bộ Text Styles có trong bản vẽ
                count = 0
                for style in doc.styles:
                    style.dxf.font = ttf_file
                    count += 1
                results.append(f"Đã đồng bộ {count} Text Styles trong file sang chuẩn Unicode ({font_name}) để sửa lỗi tiếng Việt.")
            elif action_type == "add_layer":
                lname = act.get("name", "NEW_LAYER")
                if lname not in doc.layers:
                    doc.layers.add(name=lname)
                    results.append(f"Thêm layer {lname}")
            elif action_type == "add_text":
                txt = act.get("text", "Text")
                x = act.get("x", 0)
                y = act.get("y", 0)
                layer = act.get("layer", "0")
                font_name = act.get("font_name", "Arial")
                
                if layer not in doc.layers:
                    doc.layers.add(name=layer)
                
                # Áp dụng style VIETNAMESE tương ứng cho Text
                style_name = 'VIETNAMESE_TIMES' if font_name == 'Times New Roman' else 'VIETNAMESE_ARIAL'
                msp.add_text(txt, dxfattribs={'layer': layer, 'style': style_name}).set_placement((x, y))
                results.append(f"Thêm text '{txt}' tại tọa độ ({x},{y}) trên layer {layer} (Font: {font_name})")
            elif action_type == "insert_block":
                b_name = act.get("name")
                x = act.get("x", 0)
                y = act.get("y", 0)
                layer = act.get("layer", "0")
                scale = act.get("scale", 1.0)
                rotation = act.get("rotation", 0.0)
                
                # Auto-Import Block từ Thư viện Trung tâm nếu bản vẽ thiếu
                if b_name not in doc.blocks and lib_doc and b_name in lib_doc.blocks:
                    imp = importer.Importer(lib_doc, doc)
                    imp.import_block(b_name)
                    imp.finalize()
                    results.append(f"Auto-Import thành công Block '{b_name}' từ Thư viện Trung tâm.")
                
                if b_name in doc.blocks:
                    if layer not in doc.layers:
                        doc.layers.add(name=layer)
                    msp.add_blockref(b_name, (x, y), dxfattribs={
                        'layer': layer,
                        'xscale': scale,
                        'yscale': scale,
                        'rotation': rotation
                    })
                    results.append(f"Chèn Block '{b_name}' tại ({x},{y}) layer {layer}")
                else:
                    results.append(f"Lỗi: Block '{b_name}' không tồn tại trong bản vẽ và cả Thư viện Trung tâm.")
                    
        doc.saveas(file_path)
        return f"Đã làm sạch ({audit_fixes} lỗi rác được xóa) và chỉnh sửa thành công {file_path}:\n- " + "\n- ".join(results)
    except Exception as e:
        return f"Lỗi sửa CAD (.dxf): {e}"

from src.hvac_tools import calc_psychrometrics, calc_duct_size, calc_cooling_load, calc_chw_pipe_size, calc_pump_fan_power, calc_ventilation_rate
from src.elec_tools import calc_cable_size, calc_breaker_size, calc_lighting_qty
from src.plumb_tools import calc_water_pipe, calc_water_tank, calc_plumbing_pump_head
from src.ff_tools import calc_sprinkler_qty, calc_fire_pump, calc_extinguisher_qty

tools = [
    search_standards, search_web, calculate, list_directory, read_excel, write_excel, 
    read_word, write_word, read_pdf, read_cad, write_cad, edit_cad, execute_python_code, ai_block_recovery,
    calc_psychrometrics, calc_duct_size, calc_cooling_load,
    calc_chw_pipe_size, calc_pump_fan_power, calc_ventilation_rate,
    calc_cable_size, calc_breaker_size, calc_lighting_qty,
    calc_water_pipe, calc_water_tank, calc_plumbing_pump_head,
    calc_sprinkler_qty, calc_fire_pump, calc_extinguisher_qty
]
